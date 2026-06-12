from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUTPUT_DIR = Path(__file__).resolve().parent
ACCENT = "#0F7B6C"
TEXT = "#162033"
TEXT_SOFT = "#4F5B72"


RESUMES = {
    "es": {
        "file_stem": "Juan_Sebastian_Rubiano_Hoja_de_Vida_ATS_ES",
        "name": "Juan Sebastián Rubiano",
        "headline": "Estudiante de Ingeniería de Sistemas | Analítica de Datos | Desarrollo de Software | IA Aplicada",
        "contact": [
            "Email: jsebastian000124@gmail.com",
            "LinkedIn: linkedin.com/in/juansrubiano",
            "GitHub: github.com/Zhulgon",
        ],
        "sections": [
            {
                "title": "Resumen profesional",
                "bullets": [
                    "Estudiante de décimo semestre de Ingeniería de Sistemas y tecnólogo en Desarrollo de Sistemas Informáticos, enfocado en analítica de datos, SQL, Python, visualización, automatización y desarrollo de software para procesos reales.",
                    "He desarrollado proyectos en dashboards ejecutivos, software operativo, realidad aumentada y prototipos de IA aplicada, combinando criterio técnico con contexto de negocio en el sector óptico.",
                ],
            },
            {
                "title": "Educación",
                "bullets": [
                    "Universidad Tecnológica de Santander | Ingeniería de Sistemas | Décimo semestre",
                    "Universidad Tecnológica de Santander | Tecnólogo en Desarrollo de Sistemas Informáticos",
                ],
            },
            {
                "title": "Habilidades técnicas",
                "bullets": [
                    "Datos y analítica: Excel, Python, SQL, Tableau, Power BI, Pandas, NumPy, limpieza de datos, EDA, dashboards, storytelling con datos.",
                    "Bases de datos: PostgreSQL, MySQL, NoSQL, modelado relacional, consultas, normalización, CRUD, Prisma ORM.",
                    "Desarrollo de software: JavaScript, TypeScript, React, Vite, Node.js, NestJS, HTML, CSS, APIs REST, Git, GitHub, Docker.",
                    "IA y machine learning: Scikit-learn, TensorFlow, Keras, modelos supervisados, evaluación de modelos, preparación de datasets, automatización e IA aplicada a procesos.",
                    "AR y mobile: Unity, C#, Vuforia, AR Foundation, ARCore, face tracking, image target, React Native, Expo.",
                ],
            },
            {
                "title": "Proyectos destacados",
                "bullets": [
                    "Optica Suite | Desarrollé una suite web para la gestión comercial, clínica y logística de ópticas, con módulos de pacientes, citas, ventas, historias clínicas, laboratorio, caja, auditoría y reportes. Stack: NestJS, Prisma, PostgreSQL, React, Vite, TypeScript y Docker.",
                    "HR Dashboard en Tableau | Construí un dashboard ejecutivo y detallado para analítica de recursos humanos con métricas de headcount, contrataciones, terminaciones y distribución por áreas, acompañado de dataset y documentación formal.",
                    "AR Monturas | Desarrollé una aplicación móvil en Unity para Android con catálogo 3D, image targets y face tracking experimental para pruebas de monturas en ópticas. Stack: Unity, C#, Vuforia, AR Foundation y ARCore.",
                    "App Finanzas + Hábitos | Construí una app móvil orientada a hábitos, finanzas y progreso gamificado con Expo, React Native, TypeScript, Zustand y persistencia local, pensada como producto real con roadmap y pruebas.",
                    "Data Lake Practice | Proyecto en desarrollo para practicar arquitectura de datos, ETL y organización por capas raw, processed y consumption-ready orientadas a analítica.",
                ],
            },
            {
                "title": "Experiencia relevante",
                "bullets": [
                    "Propietario y administrador de óptica | 2 años | Gestioné inventario, proveedores, flujo de caja, atención al cliente y operación comercial, identificando necesidades reales que después transformé en requerimientos y soluciones de software.",
                ],
            },
            {
                "title": "Certificaciones",
                "bullets": [
                    "Google Data Analytics y Google Advanced Data Analytics.",
                    "EF SET English Certificate 62/100 | C1 Advanced.",
                    "PCAP: Programming Essentials in Python | Cisco Networking Academy.",
                    "Ruta IBM en IA, ciberseguridad, tecnologías emergentes, mentalidad digital y empleabilidad con IA.",
                ],
            },
            {
                "title": "Idiomas",
                "bullets": [
                    "Español nativo.",
                    "Inglés C1 | EF SET.",
                ],
            },
        ],
    },
    "en": {
        "file_stem": "Juan_Sebastian_Rubiano_ATS_Resume_EN",
        "name": "Juan Sebastian Rubiano",
        "headline": "Systems Engineering Student | Data Analytics | Software Development | Applied AI",
        "contact": [
            "Email: jsebastian000124@gmail.com",
            "LinkedIn: linkedin.com/in/juansrubiano",
            "GitHub: github.com/Zhulgon",
        ],
        "sections": [
            {
                "title": "Professional summary",
                "bullets": [
                    "Tenth-semester Systems Engineering student and Software Development Technologist focused on data analytics, SQL, Python, visualization, automation and software development for real operational workflows.",
                    "I have built projects in executive dashboards, operational software, augmented reality and applied AI prototypes, combining technical execution with real business context from the optical sector.",
                ],
            },
            {
                "title": "Education",
                "bullets": [
                    "Universidad Tecnológica de Santander | Systems Engineering | Tenth semester",
                    "Universidad Tecnológica de Santander | Software Development Technologist",
                ],
            },
            {
                "title": "Technical skills",
                "bullets": [
                    "Data and analytics: Excel, Python, SQL, Tableau, Power BI, Pandas, NumPy, data cleaning, EDA, dashboards, data storytelling.",
                    "Databases: PostgreSQL, MySQL, NoSQL, relational modeling, queries, normalization, CRUD, Prisma ORM.",
                    "Software development: JavaScript, TypeScript, React, Vite, Node.js, NestJS, HTML, CSS, REST APIs, Git, GitHub, Docker.",
                    "AI and machine learning: Scikit-learn, TensorFlow, Keras, supervised models, model evaluation, dataset preparation, automation and AI applied to processes.",
                    "AR and mobile: Unity, C#, Vuforia, AR Foundation, ARCore, face tracking, image targets, React Native, Expo.",
                ],
            },
            {
                "title": "Selected projects",
                "bullets": [
                    "Optica Suite | Built a web suite for commercial, clinical and logistics workflows in optical businesses, including patients, appointments, sales, clinical records, lab orders, cash management, audit logs and reporting. Stack: NestJS, Prisma, PostgreSQL, React, Vite, TypeScript and Docker.",
                    "HR Dashboard in Tableau | Built executive and detailed HR analytics dashboards covering headcount, hiring, terminations and department distribution, supported by a structured dataset and formal documentation.",
                    "AR Monturas | Developed a Unity Android application with 3D catalog exploration, image target recognition and experimental face tracking for optical frame try-on experiences. Stack: Unity, C#, Vuforia, AR Foundation and ARCore.",
                    "Habits + Finance App | Built a product-oriented mobile app for habits, personal finance and gamified progress using Expo, React Native, TypeScript, Zustand and local persistence.",
                    "Data Lake Practice | Ongoing project to practice data architecture, ETL and layered organization using raw, processed and consumption-ready data structures.",
                ],
            },
            {
                "title": "Relevant experience",
                "bullets": [
                    "Optical business owner and administrator | 2 years | Managed inventory, suppliers, cash flow, customer service and commercial operations, identifying real workflow needs that later became software requirements and solutions.",
                ],
            },
            {
                "title": "Certifications",
                "bullets": [
                    "Google Data Analytics and Google Advanced Data Analytics.",
                    "EF SET English Certificate 62/100 | C1 Advanced.",
                    "PCAP: Programming Essentials in Python | Cisco Networking Academy.",
                    "IBM pathway in AI, cybersecurity, emerging technologies, digital mindset and AI-assisted employability.",
                ],
            },
            {
                "title": "Languages",
                "bullets": [
                    "Spanish | Native.",
                    "English | C1 | EF SET.",
                ],
            },
        ],
    },
    "es_data": {
        "file_stem": "Juan_Sebastian_Rubiano_Hoja_de_Vida_Datos_ATS_ES",
        "name": "Juan Sebastián Rubiano",
        "headline": "Estudiante de Ingeniería de Sistemas | Analítica de Datos | BI | Machine Learning | Software como Complemento",
        "contact": [
            "Email: jsebastian000124@gmail.com",
            "LinkedIn: linkedin.com/in/juansrubiano",
            "GitHub: github.com/Zhulgon",
        ],
        "sections": [
            {
                "title": "Resumen profesional",
                "bullets": [
                    "Estudiante de décimo semestre de Ingeniería de Sistemas y tecnólogo en Desarrollo de Sistemas Informáticos, orientado a analítica de datos, SQL, Python, Excel, visualización y machine learning aplicado a problemas reales.",
                    "Mi perfil se enfoca en transformar información en dashboards, análisis, criterios de decisión y automatización. El desarrollo de software actúa como complemento para convertir soluciones analíticas en productos funcionales.",
                ],
            },
            {
                "title": "Educación",
                "bullets": [
                    "Universidad Tecnológica de Santander | Ingeniería de Sistemas | Décimo semestre",
                    "Universidad Tecnológica de Santander | Tecnólogo en Desarrollo de Sistemas Informáticos",
                ],
            },
            {
                "title": "Habilidades técnicas",
                "bullets": [
                    "Datos y BI: Excel, SQL, Python, Tableau, Power BI, dashboards, storytelling con datos, análisis exploratorio de datos, limpieza de datos.",
                    "Bases de datos: PostgreSQL, MySQL, NoSQL, modelado relacional, consultas, normalización, CRUD, práctica en arquitectura de datos y data lake.",
                    "Machine learning e IA: Pandas, NumPy, Scikit-learn, TensorFlow, Keras, modelos supervisados, redes neuronales básicas, CNN, evaluación de modelos, preparación de datasets.",
                    "Complemento en software: JavaScript, TypeScript, React, Node.js, NestJS, APIs REST, Git, GitHub y Docker para llevar análisis y automatización a soluciones utilizables.",
                ],
            },
            {
                "title": "Proyectos relevantes para datos",
                "bullets": [
                    "HR Dashboard en Tableau | Construí un dashboard ejecutivo y detallado para analítica de recursos humanos con métricas de headcount, contrataciones, terminaciones y distribución por áreas, acompañado de dataset, mockups y documentación formal.",
                    "Data Lake Practice | Proyecto en desarrollo para practicar organización por capas raw, processed y consumption-ready, criterios de arquitectura de datos y flujo ETL orientado a analítica.",
                    "IA y análisis de datos | He trabajado ejercicios y prototipos con datasets de texto, audio y video para fortalecer limpieza de datos, EDA, modelos supervisados, redes neuronales básicas y experimentación con CNN.",
                    "Optica Suite | Aunque es un proyecto de software, me permitió trabajar con estructura de datos operativos, trazabilidad, reportes y necesidades reales de información dentro de un negocio óptico.",
                ],
            },
            {
                "title": "Experiencia relevante",
                "bullets": [
                    "Propietario y administrador de óptica | 2 años | Gestioné inventario, proveedores, flujo de caja, atención al cliente y operación comercial, desarrollando criterio de negocio y entendimiento de datos operativos, reportes y necesidades reales de seguimiento.",
                ],
            },
            {
                "title": "Certificaciones",
                "bullets": [
                    "Google Data Analytics y Google Advanced Data Analytics.",
                    "EF SET English Certificate 62/100 | C1 Advanced.",
                    "PCAP: Programming Essentials in Python | Cisco Networking Academy.",
                    "Ruta IBM en IA, ciberseguridad, tecnologías emergentes, mentalidad digital y empleabilidad con IA.",
                ],
            },
            {
                "title": "Idiomas",
                "bullets": [
                    "Español nativo.",
                    "Inglés C1 | EF SET.",
                ],
            },
        ],
    },
    "en_data": {
        "file_stem": "Juan_Sebastian_Rubiano_Data_Focused_ATS_Resume_EN",
        "name": "Juan Sebastian Rubiano",
        "headline": "Systems Engineering Student | Data Analytics | BI | Machine Learning | Software as a Complement",
        "contact": [
            "Email: jsebastian000124@gmail.com",
            "LinkedIn: linkedin.com/in/juansrubiano",
            "GitHub: github.com/Zhulgon",
        ],
        "sections": [
            {
                "title": "Professional summary",
                "bullets": [
                    "Tenth-semester Systems Engineering student and Software Development Technologist focused on data analytics, SQL, Python, Excel, visualization and machine learning applied to real problems.",
                    "My profile is centered on turning information into dashboards, analysis, decision support and automation. Software development works as a complement that helps transform analytical work into usable solutions.",
                ],
            },
            {
                "title": "Education",
                "bullets": [
                    "Universidad Tecnológica de Santander | Systems Engineering | Tenth semester",
                    "Universidad Tecnológica de Santander | Software Development Technologist",
                ],
            },
            {
                "title": "Technical skills",
                "bullets": [
                    "Data and BI: Excel, SQL, Python, Tableau, Power BI, dashboards, data storytelling, exploratory data analysis, data cleaning.",
                    "Databases: PostgreSQL, MySQL, NoSQL, relational modeling, queries, normalization, CRUD, practice in data architecture and data lake structure.",
                    "Machine learning and AI: Pandas, NumPy, Scikit-learn, TensorFlow, Keras, supervised models, basic neural networks, CNN experimentation, model evaluation, dataset preparation.",
                    "Software as a complement: JavaScript, TypeScript, React, Node.js, NestJS, REST APIs, Git, GitHub and Docker to turn analysis and automation into usable products.",
                ],
            },
            {
                "title": "Data-relevant projects",
                "bullets": [
                    "HR Dashboard in Tableau | Built executive and detailed HR analytics dashboards with metrics for headcount, hiring, terminations and department distribution, supported by a structured dataset, mockups and documentation.",
                    "Data Lake Practice | Ongoing project to practice raw, processed and consumption-ready layering, data architecture criteria and ETL-oriented thinking for analytics.",
                    "AI and data analysis | Worked on exercises and prototypes with text, audio and video datasets to strengthen data cleaning, EDA, supervised models, basic neural networks and CNN experimentation.",
                    "Optica Suite | Although primarily a software project, it involved working with operational data structures, traceability, reporting and real information needs in an optical business.",
                ],
            },
            {
                "title": "Relevant experience",
                "bullets": [
                    "Optical business owner and administrator | 2 years | Managed inventory, suppliers, cash flow, customer service and commercial operations, developing strong business awareness and understanding of operational data, reporting and real tracking needs.",
                ],
            },
            {
                "title": "Certifications",
                "bullets": [
                    "Google Data Analytics and Google Advanced Data Analytics.",
                    "EF SET English Certificate 62/100 | C1 Advanced.",
                    "PCAP: Programming Essentials in Python | Cisco Networking Academy.",
                    "IBM pathway in AI, cybersecurity, emerging technologies, digital mindset and AI-assisted employability.",
                ],
            },
            {
                "title": "Languages",
                "bullets": [
                    "Spanish | Native.",
                    "English | C1 | EF SET.",
                ],
            },
        ],
    },
}


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.2)


def add_docx_paragraph(document: Document, text: str, style_name: str = "Normal", *, bold=False, color=None, align=None, space_after=0) -> None:
    paragraph = document.add_paragraph(style=style_name)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(space_after)


def build_docx(resume: dict) -> None:
    document = Document()
    set_doc_defaults(document)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(resume["name"])
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(17)
    name.paragraph_format.space_after = Pt(2)

    headline = document.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = headline.add_run(resume["headline"])
    run.font.size = Pt(10.5)
    run.bold = True
    headline.paragraph_format.space_after = Pt(4)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(" | ".join(resume["contact"]))
    contact_run.font.size = Pt(9.4)
    contact.paragraph_format.space_after = Pt(8)

    for section in resume["sections"]:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run(section["title"].upper())
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10.2)

        for bullet in section["bullets"]:
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)
            paragraph.paragraph_format.space_after = Pt(1)
            bullet_run = paragraph.add_run(f"• {bullet}")
            bullet_run.font.size = Pt(10.1)

    document.save(OUTPUT_DIR / f"{resume['file_stem']}.docx")


def build_pdf(resume: dict) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=20,
        alignment=TA_CENTER,
        textColor=HexColor(TEXT),
        spaceAfter=3,
    )
    subtitle_style = ParagraphStyle(
        "ResumeSubtitle",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        alignment=TA_CENTER,
        textColor=HexColor(TEXT),
        spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "ResumeContact",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.2,
        leading=11.2,
        alignment=TA_CENTER,
        textColor=HexColor(TEXT_SOFT),
        spaceAfter=7,
    )
    section_style = ParagraphStyle(
        "ResumeSection",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=10.2,
        leading=12,
        alignment=TA_LEFT,
        textColor=HexColor(TEXT),
        spaceBefore=5,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.7,
        leading=11.6,
        alignment=TA_LEFT,
        textColor=HexColor(TEXT),
        leftIndent=12,
        firstLineIndent=-8,
        spaceAfter=1,
    )

    story = [
        Paragraph(resume["name"], title_style),
        Paragraph(resume["headline"], subtitle_style),
        Paragraph(" | ".join(resume["contact"]), contact_style),
    ]

    for section in resume["sections"]:
        story.append(Paragraph(section["title"].upper(), section_style))
        for bullet in section["bullets"]:
            story.append(Paragraph(f"• {bullet}", bullet_style))
        story.append(Spacer(1, 2))

    doc = SimpleDocTemplate(
        str(OUTPUT_DIR / f"{resume['file_stem']}.pdf"),
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    doc.build(story)


def build_markdown(resume: dict) -> None:
    lines = [
        f"# {resume['name']}",
        "",
        resume["headline"],
        "",
        " | ".join(resume["contact"]),
        "",
    ]

    for section in resume["sections"]:
        lines.append(f"## {section['title']}")
        lines.append("")
        for bullet in section["bullets"]:
            lines.append(f"- {bullet}")
        lines.append("")

    (OUTPUT_DIR / f"{resume['file_stem']}.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for resume in RESUMES.values():
        build_docx(resume)
        build_pdf(resume)
        build_markdown(resume)


if __name__ == "__main__":
    main()
